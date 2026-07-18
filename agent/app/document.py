from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .models import CandidateProfile, ResumeDraft


def _heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run(text)


def build_docx(profile: CandidateProfile, resume: ResumeDraft, destination: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    name_run = name.add_run(profile.name.upper())
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_items = [
        profile.location,
        profile.phone,
        profile.email,
        profile.linkedin,
        profile.github_repo,
        profile.portfolio,
    ]
    contact = document.add_paragraph(" | ".join(item for item in contact_items if item))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)

    if resume.professional_summary:
        _heading(document, "Professional Summary")
        document.add_paragraph(resume.professional_summary)

    if resume.skills:
        _heading(document, "Skills")
        document.add_paragraph(" • ".join(resume.skills))

    if resume.experience:
        _heading(document, "Experience")
        for entry in resume.experience:
            line = document.add_paragraph()
            line.paragraph_format.space_before = Pt(3)
            line.paragraph_format.space_after = Pt(0)
            run = line.add_run(f"{entry.title} — {entry.organization}")
            run.bold = True
            dates = " – ".join(item for item in [entry.start_date, entry.end_date] if item)
            details = " | ".join(item for item in [entry.location, dates] if item)
            if details:
                line.add_run(f"  |  {details}")
            for bullet in entry.bullets:
                _bullet(document, bullet)

    if resume.projects:
        _heading(document, "Projects")
        for project in resume.projects:
            line = document.add_paragraph()
            line.paragraph_format.space_before = Pt(3)
            line.paragraph_format.space_after = Pt(0)
            title = project.name + (f" | {project.link}" if project.link else "")
            run = line.add_run(title)
            run.bold = True
            if project.technologies:
                line.add_run(f" — {', '.join(project.technologies)}")
            for bullet in project.bullets:
                _bullet(document, bullet)

    if resume.education:
        _heading(document, "Education")
        for entry in resume.education:
            line = document.add_paragraph()
            line.paragraph_format.space_before = Pt(3)
            line.paragraph_format.space_after = Pt(0)
            run = line.add_run(f"{entry.credential} — {entry.institution}")
            run.bold = True
            details = " | ".join(item for item in [entry.location, entry.graduation_date] if item)
            if details:
                line.add_run(f"  |  {details}")
            for detail in entry.details:
                _bullet(document, detail)

    if resume.certifications:
        _heading(document, "Certifications")
        for certification in resume.certifications:
            _bullet(document, certification)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
