import os
from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException
from fastapi.testclient import TestClient
from sqlalchemy import create_engine

os.environ.setdefault("JWT_SECRET", "test-secret")

import main  # noqa: E402
from main import VisibleTextParser, extract_document, validate_public_url  # noqa: E402


def test_document_and_job_input_guards(tmp_path: Path):
    path = tmp_path / "background.docx"
    document = Document()
    document.add_paragraph("Built a FastAPI service for resume generation.")
    document.save(path)
    assert "FastAPI service" in extract_document(path)

    parser = VisibleTextParser()
    parser.feed("<main>Python engineer<style>ignore me</style><script>ignore()</script></main>")
    assert parser.parts == ["Python engineer"]

    with pytest.raises(HTTPException):
        validate_public_url("http://127.0.0.1/internal")


def test_authenticated_generation_flow(tmp_path: Path, monkeypatch):
    main.engine = create_engine(
        f"sqlite:///{tmp_path / 'test.db'}",
        connect_args={"check_same_thread": False},
    )
    main.UPLOAD_DIR = tmp_path / "uploads"
    main.CAREER_DIR = tmp_path / "career"
    main.GENERATED_DIR = tmp_path / "generated"
    for directory in [main.UPLOAD_DIR, main.CAREER_DIR, main.GENERATED_DIR]:
        directory.mkdir()
    main.Base.metadata.create_all(main.engine)
    monkeypatch.setenv("AGENT_API_KEY", "internal-test-key")

    generated = tmp_path / "agent.docx"
    Document().save(generated)

    class FakeResponse:
        def __init__(self, data=None, content=b""):
            self._data, self.content, self.is_success = data, content, True

        def raise_for_status(self):
            return None

        def json(self):
            return self._data

    class FakeAgentClient:
        def __init__(self, *args, **kwargs):
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            pass

        async def post(self, url, json, headers):
            assert "FastAPI" in json["profile"]["resume_text"]
            return FakeResponse({"compatibility_score": 90, "docx_url": "/resume-generations/agent/resume.docx"})

        async def get(self, url, headers):
            return FakeResponse(content=generated.read_bytes())

    monkeypatch.setattr(main.httpx, "AsyncClient", FakeAgentClient)
    client = TestClient(main.app)
    auth = client.post("/auth/register", json={"name": "Alex", "email": "alex@example.com", "password": "password123"}).json()
    headers = {"Authorization": f"Bearer {auth['access_token']}"}

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Built FastAPI applications for clients.")
    document.save(source)
    with source.open("rb") as upload:
        assert client.post("/account/documents/resume", headers=headers, files={"document": ("resume.docx", upload)}).status_code == 200

    response = client.post(
        "/resume-generations",
        headers=headers,
        json={"job_description": "Seeking a Python and FastAPI engineer to build reliable web services for our growing product team."},
    )
    assert response.status_code == 200
    download = client.get(response.json()["docx_url"], headers=headers)
    assert download.status_code == 200
    assert download.content.startswith(b"PK")
